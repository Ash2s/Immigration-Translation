import pytest
import tempfile
import os
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from lxml import etree
from docx.oxml.ns import qn
from app.services.document_parser import DocumentParser


def test_read_docx_paragraphs():
    parser = DocumentParser()
    doc = DocxDocument()
    doc.add_paragraph("Hello, World!")
    fd, path = tempfile.mkstemp(suffix='.docx')
    os.close(fd)
    try:
        doc.save(path)
        loaded = parser.read_document(path)
        paragraphs = parser.extract_paragraphs(loaded)
        assert len(paragraphs) >= 1
        assert paragraphs[0]["text"] == "Hello, World!"
    finally:
        os.unlink(path)


def test_preserve_bold():
    parser = DocumentParser()
    doc = DocxDocument()
    p = doc.add_paragraph()
    run = p.add_run("Bold text")
    run.font.bold = True
    fd, path = tempfile.mkstemp(suffix='.docx')
    os.close(fd)
    try:
        doc.save(path)
        loaded = parser.read_document(path)
        paragraphs = parser.extract_paragraphs(loaded)
        assert len(paragraphs[0]["runs"]) >= 1
        assert paragraphs[0]["runs"][0]["bold"] is True
    finally:
        os.unlink(path)


def test_preserve_italic():
    parser = DocumentParser()
    doc = DocxDocument()
    p = doc.add_paragraph()
    run = p.add_run("Italic text")
    run.font.italic = True
    fd, path = tempfile.mkstemp(suffix='.docx')
    os.close(fd)
    try:
        doc.save(path)
        loaded = parser.read_document(path)
        paragraphs = parser.extract_paragraphs(loaded)
        assert paragraphs[0]["runs"][0]["italic"] is True
    finally:
        os.unlink(path)


def test_preserve_alignment():
    parser = DocumentParser()
    doc = DocxDocument()
    p = doc.add_paragraph("Right aligned")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fd, path = tempfile.mkstemp(suffix='.docx')
    os.close(fd)
    try:
        doc.save(path)
        loaded = parser.read_document(path)
        paragraphs = parser.extract_paragraphs(loaded)
        assert paragraphs[0]["alignment"] == WD_ALIGN_PARAGRAPH.RIGHT.value
    finally:
        os.unlink(path)


def test_clear_shading():
    """Verify that paragraph-level background shading is removed by clear_background_shading."""
    parser = DocumentParser()
    doc = DocxDocument()
    p = doc.add_paragraph("Shaded text")

    # Inject paragraph-level shading via XML
    pPr = p._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = etree.SubElement(p._element, qn('w:pPr'))
    shd = etree.SubElement(pPr, qn('w:shd'))
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'D9D9D9')

    fd, path = tempfile.mkstemp(suffix='.docx')
    os.close(fd)
    try:
        doc.save(path)
        loaded = parser.read_document(path)
        parser.clear_background_shading(loaded)

        for para in loaded.paragraphs:
            pPr = para._element.find(qn('w:pPr'))
            if pPr is not None:
                shd = pPr.find(qn('w:shd'))
                assert shd is None, "Paragraph-level shading should be removed"
    finally:
        os.unlink(path)


def test_output_format_times_new_roman():
    """After apply_formatting, verify font is Times New Roman 12pt with bold preserved."""
    parser = DocumentParser()
    doc = DocxDocument()
    p = doc.add_paragraph()
    run = p.add_run("Original bold text")
    run.font.bold = True
    run.font.name = "SimSun"
    run.font.size = Pt(14)

    fd, path = tempfile.mkstemp(suffix='.docx')
    os.close(fd)
    try:
        doc.save(path)
        loaded = parser.read_document(path)
        paragraphs = parser.extract_paragraphs(loaded)
        runs_data = paragraphs[0]["runs"]
        original_para = loaded.paragraphs[0]

        parser.apply_formatting(original_para, runs_data, "Translated text")

        # Save and reload to read back the formatting
        loaded.save(path)
        reloaded = parser.read_document(path)

        para = reloaded.paragraphs[0]
        assert para.text == "Translated text"
        assert len(para.runs) >= 1
        assert para.runs[0].font.name == "Times New Roman"
        assert para.runs[0].font.size == Pt(12)
        assert para.runs[0].font.bold is True
    finally:
        os.unlink(path)


def test_apply_per_run_formatting_adds_space_between_runs():
    """When per-run translations would merge without spaces (e.g. Chinese
    character fragments translated independently), apply_per_run_formatting
    should insert a space between adjacent alphanumeric runs."""
    parser = DocumentParser()
    doc = DocxDocument()
    p = doc.add_paragraph()

    # Simulate two runs that were originally Chinese characters "例" "如"
    r1 = p.add_run("Example")
    r2 = p.add_run("For example")

    runs_data = [
        {"text": "例", "bold": False, "italic": False, "underline": False,
         "font_name": None, "font_size": None, "highlight": None, "color": None},
        {"text": "如", "bold": False, "italic": False, "underline": False,
         "font_name": None, "font_size": None, "highlight": None, "color": None},
    ]
    translated_runs = ["Example", "For example"]

    parser.apply_per_run_formatting(p, runs_data, translated_runs)

    # The two runs should now have a space between them
    assert p.text == "Example For example", (
        f"Expected 'Example For example', got: {repr(p.text)}"
    )

