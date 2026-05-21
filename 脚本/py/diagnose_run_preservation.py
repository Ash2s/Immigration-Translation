"""
Diagnose run-structure preservation across direct-call vs thread-pool-call.

Tests whether `DocumentParser.apply_per_run_formatting` preserves per-run text
and formatting when called in different thread contexts (main thread == direct,
vs thread pool == web flow).

Usage:
    python 脚本/py/diagnose_run_preservation.py [file_id]

If no file_id is given, uses "test_live" which must exist in data/uploads/.
"""
import sys
import os
import asyncio
import tempfile

# Ensure project root is on sys.path
_PROJ = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
sys.path.insert(0, _PROJ)

from app.services.document_parser import DocumentParser
from docx import Document as DocxDocument
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

UPLOAD_DIR = os.path.join(_PROJ, "data", "uploads")


def _count_runs_per_para(path: str) -> list[int]:
    """Return list of run-count per body paragraph."""
    doc = DocxDocument(path)
    return [len(p.runs) for p in doc.paragraphs]


def _has_multi_run(path: str) -> bool:
    """Return True if any body paragraph has > 1 run."""
    return any(c > 1 for c in _count_runs_per_para(path))


def _print_multi_run_details(path: str, label: str):
    """Print paragraphs where run count > 1, showing run texts."""
    doc = DocxDocument(path)
    print(f"\n  [{label}] Multi-run paragraphs:")
    found = False
    for i, p in enumerate(doc.paragraphs):
        if len(p.runs) > 1:
            found = True
            print(f"    Para {i}: {len(p.runs)} runs")
            for j, r in enumerate(p.runs):
                hl = r.font.highlight_color
                ul = r.font.underline
                hl_str = hl.name if hl else "-"
                ul_str = str(ul) if ul else "-"
                txt = r.text[:50].replace("\n", "\\n")
                print(f"      Run {j}: hl={hl_str} ul={ul_str} text={txt!r}")
    if not found:
        print("    (none)")


def _check_multi_run_in_outputs(outputs: dict[str, str]):
    """Compare multi-run preservation across all output variants."""
    print("\n========== RUN STRUCTURE COMPARISON ==========")
    source_counts = _count_runs_per_para(outputs["source"])
    print(f"\nSource run counts: {source_counts}")
    print(f"Source multi-run paragraphs: {sum(1 for c in source_counts if c > 1)}")

    for label, path in outputs.items():
        if label == "source":
            continue
        counts = _count_runs_per_para(path)
        if counts == source_counts:
            print(f"  ✓ [{label}] Run structure matches source")
        else:
            mismatches = [(i, s, c) for i, (s, c) in enumerate(zip(source_counts, counts)) if s != c]
            print(f"  ✗ [{label}] Run structure DIFFERS:")
            for i, s, c in mismatches:
                print(f"      Para {i}: source={s} runs, output={c} runs")


def diagnose_file(file_id: str):
    """Run apply_per_run_formatting in 4 scenarios and compare outputs."""
    src_path = os.path.join(UPLOAD_DIR, f"{file_id}.docx")
    if not os.path.exists(src_path):
        print(f"ERROR: source file not found: {src_path}")
        return

    # 1. Open source doc and extract paragraph data
    parser = DocumentParser()
    doc = parser.read_document(src_path)
    paragraphs_data = parser.extract_paragraphs(doc)

    # Generate mock translated runs (simulate what _translate_paragraphs_sync does)
    mock_translated = []
    for pd in paragraphs_data:
        mock_runs = []
        for rd in pd["runs"]:
            if not rd["text"].strip():
                mock_runs.append(rd["text"])
            else:
                mock_runs.append(f"[TRANS_{len(mock_runs)}]")
        mock_translated.append(mock_runs)

    # Track output paths
    outputs = {"source": src_path}

    # ── A) Direct call (main thread) ──
    doc_a = parser.read_document(src_path)
    for idx, pd in enumerate(paragraphs_data):
        if idx < len(doc_a.paragraphs):
            parser.apply_per_run_formatting(
                doc_a.paragraphs[idx], pd["runs"], mock_translated[idx]
            )
    path_a = os.path.join(UPLOAD_DIR, f"diagnose_A_{file_id}.docx")
    parser.save_document(doc_a, path_a)
    outputs["A-direct"] = path_a
    _print_multi_run_details(path_a, "A-direct (main thread)")

    # ── B) Direct call + real glossary (just test glossary doesn't affect format) ──
    doc_b = parser.read_document(src_path)
    for idx, pd in enumerate(paragraphs_data):
        if idx < len(doc_b.paragraphs):
            parser.apply_per_run_formatting(
                doc_b.paragraphs[idx], pd["runs"], mock_translated[idx]
            )
    path_b = os.path.join(UPLOAD_DIR, f"diagnose_B_{file_id}.docx")
    parser.save_document(doc_b, path_b)
    outputs["B-direct+glossary"] = path_b

    # ── C) run_in_executor (thread pool) ──
    async def _run_in_thread(executor=None):
        loop = asyncio.get_event_loop()
        doc_c = await loop.run_in_executor(executor, parser.read_document, src_path)
        paras_c = parser.extract_paragraphs(doc_c)
        for idx, pd in enumerate(paras_c):
            if idx < len(doc_c.paragraphs):
                parser.apply_per_run_formatting(
                    doc_c.paragraphs[idx], pd["runs"], mock_translated[idx]
                )
        path_c = os.path.join(UPLOAD_DIR, f"diagnose_C_{file_id}.docx")
        await loop.run_in_executor(executor, parser.save_document, doc_c, path_c)
        return path_c

    path_c = asyncio.run(_run_in_thread())
    outputs["C-threadpool"] = path_c
    _print_multi_run_details(path_c, "C-threadpool")

    # ── D) Thread pool + save in executor (closest to web flow) ──
    async def _run_full_in_thread():
        loop = asyncio.get_event_loop()
        doc_d = await loop.run_in_executor(None, parser.read_document, src_path)
        paras_d = parser.extract_paragraphs(doc_d)
        def _apply():
            for idx, pd in enumerate(paras_d):
                if idx < len(doc_d.paragraphs):
                    parser.apply_per_run_formatting(
                        doc_d.paragraphs[idx], pd["runs"], mock_translated[idx]
                    )
            path_d = os.path.join(UPLOAD_DIR, f"diagnose_D_{file_id}.docx")
            parser.save_document(doc_d, path_d)
            return path_d
        return await loop.run_in_executor(None, _apply)

    path_d = asyncio.run(_run_full_in_thread())
    outputs["D-fullthread"] = path_d
    _print_multi_run_details(path_d, "D-fullthread (full in thread)")

    # ── Comparison ──
    _check_multi_run_in_outputs(outputs)

    print("\n========== SUMMARY ==========")
    # Clean up temp files (enable if desired: set CLEANUP=1)
    if os.environ.get("CLEANUP"):
        for label, path in outputs.items():
            if label != "source" and os.path.exists(path):
                os.remove(path)
                print(f"  Removed: {path}")
    else:
        print("  Files saved to data/uploads/ for inspection:")
        for label, path in outputs.items():
            if label != "source":
                print(f"    {label}: {os.path.basename(path)}")


if __name__ == "__main__":
    file_id = sys.argv[1] if len(sys.argv) > 1 else "test_live"
    diagnose_file(file_id)
