from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from typing import Optional
from lxml import etree
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re


class DocumentParser:
    """
    Service for reading .docx documents, extracting paragraph and run-level
    formatting, applying formatting to translated text, clearing background
    shading, and saving results.

    Follows the same service pattern as GlossaryService.
    """

    # OOXML namespace for the w: prefix
    _NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def read_document(self, path: str) -> DocxDocument:
        """Open a .docx file and return a python-docx Document object."""
        return DocxDocument(path)

    def extract_paragraphs(self, doc: DocxDocument) -> list[dict]:
        """
        Extract every paragraph from *doc* along with full run-level
        formatting information.

        Returns a list of dicts.  Each dict has:
          - ``text``      — full paragraph text
          - ``alignment`` — integer value of WD_ALIGN_PARAGRAPH or None
          - ``runs``      — list of run dicts (see _extract_run_data)
        """
        paragraphs = []
        for para in doc.paragraphs:
            runs_data = [self._extract_run_data(run) for run in para.runs]
            paragraphs.append({
                "text": para.text,
                "alignment": para.alignment.value
                              if para.alignment is not None else None,
                "runs": runs_data,
            })
        return paragraphs

    def extract_table_cells(self, doc: DocxDocument) -> list[dict]:
        """
        Extract text and run formatting from every non-empty table cell
        paragraph in every table.

        Returns a list of dicts compatible with ``_translate_paragraphs``,
        plus a ``paragraph`` key holding the python-docx Paragraph object
        for write-back.
        """
        cells = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if not para.text.strip():
                            continue
                        runs_data = [
                            self._extract_run_data(run) for run in para.runs
                        ]
                        cells.append({
                            "text": para.text,
                            "alignment": (
                                para.alignment.value
                                if para.alignment is not None else None
                            ),
                            "runs": runs_data,
                            "paragraph": para,
                        })
        return cells

    def extract_textbox_paragraphs(self, doc: DocxDocument) -> list[dict]:
        """
        Extract text and formatting from every text box (``<w:txbxContent>``)
        paragraph in the document. Text boxes are not accessible via
        ``doc.paragraphs`` or ``doc.tables``, so this method works at the
        XML level.

        Returns a list of dicts with ``text``, ``runs``, and ``element``
        (the ``w:p`` XML element, for write-back).
        """
        paragraphs = []
        for txbx in doc.element.findall(f'.//{qn("w:txbxContent")}'):
            for p_elem in txbx.findall(qn('w:p')):
                text = self._get_element_text(p_elem)
                if not text.strip():
                    continue
                runs_data = []
                for r in p_elem.findall(qn('w:r')):
                    t_el = r.find(qn('w:t'))
                    r_text = t_el.text if t_el is not None and t_el.text else ''
                    rPr = r.find(qn('w:rPr'))
                    runs_data.append({
                        "text": r_text,
                        "bold": self._rpr_has_bold(rPr),
                        "italic": self._rpr_has_italic(rPr),
                        "underline": False,
                        "font_name": None,
                        "font_size": None,
                        "highlight": None,
                        "color": None,
                    })
                paragraphs.append({
                    "text": text,
                    "runs": runs_data,
                    "element": p_elem,
                })
        return paragraphs

    @staticmethod
    def _get_element_text(p_elem) -> str:
        """Extract concatenated text from a ``w:p`` XML element."""
        parts = []
        for r in p_elem.findall(qn('w:r')):
            t_el = r.find(qn('w:t'))
            if t_el is not None and t_el.text:
                parts.append(t_el.text)
            for br in r.findall(qn('w:br')):
                parts.append('\n')
        return ''.join(parts)

    @staticmethod
    def _rpr_has_bold(rPr) -> bool:
        """Check if an ``rPr`` element has bold enabled."""
        if rPr is None:
            return False
        b = rPr.find(qn('w:b'))
        return b is not None and b.get(qn('w:val'), '1') != '0'

    @staticmethod
    def _rpr_has_italic(rPr) -> bool:
        """Check if an ``rPr`` element has italic enabled."""
        if rPr is None:
            return False
        i = rPr.find(qn('w:i'))
        return i is not None and i.get(qn('w:val'), '1') != '0'

    @staticmethod
    def apply_textbox_formatting(p_elem, runs_data: list[dict], text: str) -> None:
        """Replace content of a text box paragraph XML element with translated text.

        Forces Times New Roman 12pt and preserves bold/italic from the
        original runs (union across all runs).
        """
        # Remove existing runs
        for r in list(p_elem.findall(qn('w:r'))):
            p_elem.remove(r)

        # Create single new run
        new_r = OxmlElement('w:r')
        new_rPr = OxmlElement('w:rPr')
        new_r.append(new_rPr)

        # Force TNR 12pt
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        new_rPr.append(rFonts)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '24')
        new_rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), '24')
        new_rPr.append(szCs)

        # Formatting union
        if any(rd.get("bold") for rd in runs_data):
            new_rPr.append(OxmlElement('w:b'))
        if any(rd.get("italic") for rd in runs_data):
            new_rPr.append(OxmlElement('w:i'))

        # Translated text
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        new_r.append(t)

        p_elem.append(new_r)

    @staticmethod
    def set_textbox_line_spacing(p_elem) -> None:
        """Set single spacing on a text box paragraph XML element."""
        pPr = p_elem.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            p_elem.insert(0, pPr)
        sp = pPr.find(qn('w:spacing'))
        if sp is None:
            sp = OxmlElement('w:spacing')
            pPr.append(sp)
        sp.set(qn('w:line'), '240')
        sp.set(qn('w:lineRule'), 'auto')
        sp.set(qn('w:before'), '0')
        sp.set(qn('w:after'), '0')

    def apply_formatting(
        self,
        para,
        runs_data: list[dict],
        text: str,
    ) -> None:
        """
        Replace the content of *para* with *text* while:
          - Setting the font to Times New Roman 12pt
          - Preserving the original paragraph alignment
          - Applying any **bold** / *italic* / underline / highlight / color
            that existed in the original runs (union across all runs).

        Parameters
        ----------
        para :
            A python-docx ``Paragraph`` object (from the same document
            that ``extract_paragraphs`` was called on).
        runs_data : list[dict]
            The list of run dicts returned by ``extract_paragraphs`` for
            this paragraph.
        text : str
            The translated text to place in the paragraph.
        """
        # 1. Find all run elements at the XML level.  We must NOT remove
        #    run elements because some of them may contain drawings or text
        #    boxes (``w:txbxContent``).  Instead we clear every ``<w:t>``
        #    child and add one new ``<w:t>`` to the first run.
        w_r_els = para._p.findall(qn('w:r'))

        if w_r_els:
            # Strip all w:t children from every run (preserving non-text
            # children such as <w:drawing>, <w:pict>, etc.)
            for r_elem in w_r_els:
                for t_elem in r_elem.findall(qn('w:t')):
                    r_elem.remove(t_elem)

            # Add a single w:t to the first run with the translated text.
            first_t = OxmlElement('w:t')
            first_t.set(qn('xml:space'), 'preserve')
            first_t.text = text
            w_r_els[0].append(first_t)

            # Apply fonts & formatting on the first run via python-docx API.
            run = para.runs[0]
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            self._apply_run_formatting(run, runs_data)
        else:
            # No runs at all – add a new one.
            new_run = para.add_run(text)
            new_run.font.name = "Times New Roman"
            new_run.font.size = Pt(12)
            self._apply_run_formatting(new_run, runs_data)

    def clear_background_shading(self, doc: DocxDocument) -> None:
        """
        Remove ``<w:shd>`` elements from every paragraph and run in the
        document, including content inside tables.  This is required by
        immigration-bureau scan specifications that forbid gray background
        shading.

        Handles both paragraph-level (``<w:pPr><w:shd>``) and run-level
        (``<w:rPr><w:shd>``) shading.
        """
        for para in doc.paragraphs:
            DocumentParser._clear_para_shading(para)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        DocumentParser._clear_para_shading(para)
        for txbx in doc.element.findall(f'.//{qn("w:txbxContent")}'):
            for p_elem in txbx.findall(qn('w:p')):
                DocumentParser._clear_shading_from_element(
                    p_elem, DocumentParser._NS + 'pPr'
                )
                for r in p_elem.findall(qn('w:r')):
                    DocumentParser._clear_shading_from_element(
                        r, DocumentParser._NS + 'rPr'
                    )

    @staticmethod
    def _clear_para_shading(para) -> None:
        """Remove background shading from a single paragraph element."""
        DocumentParser._clear_shading_from_element(
            para._element, DocumentParser._NS + 'pPr'
        )
        for run in para.runs:
            DocumentParser._clear_shading_from_element(
                run._element, DocumentParser._NS + 'rPr'
            )

    @staticmethod
    def set_line_spacing(para, double_space: bool = True) -> None:
        """Set paragraph line spacing. True = double (2x), False = single (1x)."""
        pPr = para._p.get_or_add_pPr()
        sp = pPr.find(qn('w:spacing'))
        if sp is None:
            sp = OxmlElement('w:spacing')
            pPr.append(sp)
        sp.set(qn('w:line'), '480' if double_space else '240')
        sp.set(qn('w:lineRule'), 'auto')
        if not double_space:
            sp.set(qn('w:before'), '0')
            sp.set(qn('w:after'), '0')

    def save_document(self, doc: DocxDocument, path: str) -> None:
        """Save the document to *path* and post-process fonts."""
        doc.save(path)
        DocumentParser.fix_document_fonts(path)

    @staticmethod
    def fix_document_fonts(docx_path: str) -> None:
        """Post-process a saved .docx to force Times New Roman at every level
        (theme XML → style defaults → run fonts)."""
        DocumentParser._fix_theme_xml(docx_path)
        doc = DocxDocument(docx_path)
        DocumentParser._fix_run_fonts(doc)
        DocumentParser._fix_style_fonts(doc)
        doc.save(docx_path)

    @staticmethod
    def _fix_theme_xml(docx_path: str) -> None:
        """Replace 宋体 with Times New Roman in the theme font scheme."""
        import zipfile, shutil
        temp_path = docx_path + '.tmp'
        with zipfile.ZipFile(docx_path, 'r') as zin:
            with zipfile.ZipFile(temp_path, 'w', zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    data = zin.read(item.filename)
                    if 'theme' in item.filename.lower():
                        text = data.decode('utf-8')
                        text = text.replace(
                            '<a:font script="Hans" typeface="宋体"/>',
                            '<a:font script="Hans" typeface="Times New Roman"/>'
                        )
                        data = text.encode('utf-8')
                    zout.writestr(item, data)
        shutil.move(temp_path, docx_path)

    @staticmethod
    def _fix_run_fonts(doc) -> None:
        """Remove theme font refs from all runs; set explicit TNR."""
        for para in list(doc.paragraphs):
            DocumentParser._fix_run_fonts_in_para(para)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        DocumentParser._fix_run_fonts_in_para(para)
        # Also fix font refs in text boxes
        for txbx in doc.element.findall(f'.//{qn("w:txbxContent")}'):
            for p_elem in txbx.findall(qn('w:p')):
                DocumentParser._fix_run_fonts_in_para_elem(p_elem)

    @staticmethod
    def _fix_run_fonts_in_para(para) -> None:
        """Fix font references in all runs of a single paragraph."""
        DocumentParser._fix_run_fonts_in_para_elem(para._p)

    @staticmethod
    def _fix_run_fonts_in_para_elem(p_elem) -> None:
        """Fix font references in all runs of a paragraph XML element."""
        for r in p_elem.findall(qn('w:r')):
            rPr = r.find(qn('w:rPr'))
            if rPr is None:
                rPr = etree.SubElement(r, qn('w:rPr'))
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = etree.SubElement(rPr, qn('w:rFonts'))
            for attr in list(rFonts.attrib):
                if 'theme' in attr.lower():
                    del rFonts.attrib[attr]
            for tag in ('w:ascii', 'w:eastAsia', 'w:hAnsi', 'w:cs'):
                if not rFonts.get(qn(tag)):
                    rFonts.set(qn(tag), 'Times New Roman')

    @staticmethod
    def _fix_style_fonts(doc) -> None:
        """Set docDefaults + Normal style to explicit TNR; clear theme refs."""
        styles_elem = doc.styles.element

        docDefaults = styles_elem.find(qn('w:docDefaults'))
        if docDefaults is None:
            docDefaults = etree.SubElement(styles_elem, qn('w:docDefaults'))
        rPrDefault = docDefaults.find(qn('w:rPrDefault'))
        if rPrDefault is None:
            rPrDefault = etree.SubElement(docDefaults, qn('w:rPrDefault'))
        rPr = rPrDefault.find(qn('w:rPr'))
        if rPr is None:
            rPr = etree.SubElement(rPrDefault, qn('w:rPr'))
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = etree.SubElement(rPr, qn('w:rFonts'))
        for tag in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
            rFonts.set(qn(tag), 'Times New Roman')
        for attr in list(rFonts.attrib):
            if 'theme' in attr.lower():
                del rFonts.attrib[attr]

        normal = doc.styles['Normal']
        n_rPr = normal.element.find(qn('w:rPr'))
        if n_rPr is None:
            n_rPr = etree.SubElement(normal.element, qn('w:rPr'))
        n_rFonts = n_rPr.find(qn('w:rFonts'))
        if n_rFonts is None:
            n_rFonts = etree.SubElement(n_rPr, qn('w:rFonts'))
        for tag in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
            n_rFonts.set(qn(tag), 'Times New Roman')
        for attr in list(n_rFonts.attrib):
            if 'theme' in attr.lower():
                del n_rFonts.attrib[attr]

        for style in styles_elem.findall(qn('w:style')):
            style_rPr = style.find(qn('w:rPr'))
            if style_rPr is not None:
                style_rFonts = style_rPr.find(qn('w:rFonts'))
                if style_rFonts is not None:
                    for attr in list(style_rFonts.attrib):
                        if 'theme' in attr.lower():
                            del style_rFonts.attrib[attr]

    @staticmethod
    def verify_no_cn(doc_or_path) -> list[str]:
        """Scan all paragraphs + table cells + text boxes for Chinese characters.
        Returns a list of location descriptions (empty = clean)."""
        cn_pat = re.compile(r'[一-鿿]')
        if isinstance(doc_or_path, str):
            doc = DocxDocument(doc_or_path)
        else:
            doc = doc_or_path
        issues = []
        for i, para in enumerate(doc.paragraphs):
            if cn_pat.search(para.text):
                issues.append(f"Para {i}: {para.text.strip()[:60]}")
        for ti, table in enumerate(doc.tables):
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    if cn_pat.search(cell.text):
                        issues.append(f"Table {ti}R{ri}C{ci}: {cell.text.strip()[:60]}")
        # Scan text boxes
        for tbi, txbx in enumerate(doc.element.findall(f'.//{qn("w:txbxContent")}')):
            for pi, p_elem in enumerate(txbx.findall(qn('w:p'))):
                text = DocumentParser._get_element_text(p_elem)
                if text.strip() and cn_pat.search(text):
                    issues.append(f"TextBox {tbi}P{pi}: {text.strip()[:60]}")
        return issues

    # ------------------------------------------------------------------
    # Internal helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _extract_run_data(run) -> dict:
        """
        Return a serialisable dict describing one run's formatting.

        Returns keys:
          text, bold, italic, underline, font_name, font_size (Pt),
          highlight (int | None), color (hex str | None)
        """
        font = run.font
        highlight = None
        if font.highlight_color is not None:
            highlight = font.highlight_color.value  # e.g. 7 for YELLOW

        color = None
        if font.color is not None and font.color.rgb is not None:
            color = str(font.color.rgb).upper()

        return {
            "text": run.text,
            "bold": bool(font.bold) if font.bold is not None else False,
            "italic": bool(font.italic) if font.italic is not None else False,
            "underline": (
                bool(font.underline)
                if font.underline is not None
                else False
            ),
            "font_name": font.name,
            "font_size": font.size.pt if font.size is not None else None,
            "highlight": highlight,
            "color": color,
        }

    @staticmethod
    def _apply_run_formatting(run, runs_data: list[dict]) -> None:
        """
        Inspect the union of formatting across all *runs_data* entries and
        apply the active properties to *run*.

        - **bold**: True if ANY original run was bold.
        - *italic*: True if ANY original run was italic.
        - underline: True if ANY original run was underlined.
        - highlight: taken from the first run that carried one.
        - color: taken from the first run that carried one.
        """
        # Determine which properties are "active" across all runs.
        has_bold = any(rd.get("bold") for rd in runs_data)
        has_italic = any(rd.get("italic") for rd in runs_data)
        has_underline = any(rd.get("underline") for rd in runs_data)
        has_highlight = any(
            rd.get("highlight") is not None for rd in runs_data
        )
        has_color = any(rd.get("color") is not None for rd in runs_data)

        if has_bold:
            run.font.bold = True
        if has_italic:
            run.font.italic = True
        if has_underline:
            run.font.underline = True

        if has_highlight:
            for rd in runs_data:
                h = rd.get("highlight")
                if h is not None:
                    try:
                        run.font.highlight_color = WD_COLOR_INDEX(h)
                    except (ValueError, KeyError):
                        pass
                    break

        if has_color:
            for rd in runs_data:
                c = rd.get("color")
                if c is not None:
                    try:
                        run.font.color.rgb = RGBColor.from_string(c)
                    except (ValueError, AttributeError):
                        pass
                    break

    @staticmethod
    def _clear_shading_from_element(element, container_tag: str) -> None:
        """
        Find the XML child element named *container_tag* (e.g. ``w:pPr`` or
        ``w:rPr``) inside *element* and remove any ``<w:shd>`` child from
        it.
        """
        NS = DocumentParser._NS
        container = element.find(container_tag)
        if container is not None:
            shd = container.find(NS + 'shd')
            if shd is not None:
                container.remove(shd)

    # ------------------------------------------------------------------
    # Formatting instruction framework
    # ------------------------------------------------------------------

    # Mapping of keywords → (target, action, value)
    # Used by apply_formatting_instructions() to interpret user feedback.
    _FORMAT_RULES: list[tuple[list[str], str, object]] = [
        # Line spacing
        (["单倍行距", "single spacing", "一倍行距"], "line_spacing", "single"),
        (["双倍行距", "double spacing", "两倍行距"], "line_spacing", "double"),
        (["1.5倍行距", "1.5 spacing", "一点五倍行距"], "line_spacing", "1.5"),
        # Font color
        (["红色", "red"], "font_color", "FF0000"),
        (["蓝色", "blue"], "font_color", "0000FF"),
        (["绿色", "green"], "font_color", "00FF00"),
        (["黑色", "black"], "font_color", "000000"),
        # Font name
        (["Times New Roman", "TNR"], "font_name", "Times New Roman"),
        (["宋体"], "font_name", "宋体"),
        (["黑体"], "font_name", "黑体"),
        # Bold / italic
        (["加粗", "bold"], "bold", True),
        (["斜体", "italic"], "italic", True),
        # Remove shading
        (["去除底纹", "remove shading", "clear shading", "清除底纹"], "clear_shading", True),
    ]

    @staticmethod
    def apply_formatting_instructions(doc: DocxDocument, feedback: str) -> list[str]:
        """DEPRECATED: applies formatting globally.  Use apply_targeted_formatting instead."""
        applied: list[str] = []
        if not feedback:
            return applied
        for keywords, action, value in DocumentParser._FORMAT_RULES:
            if any(kw in feedback for kw in keywords):
                DocumentParser._apply_format_action(doc, action, value)
                applied.append(f"{action} → {value}")
        return applied

    @staticmethod
    def apply_targeted_formatting(
        doc: DocxDocument,
        actions: list[dict],
    ) -> list[str]:
        """Apply formatting actions to specific paragraph indices.

        Each action dict has:
          - ``indices`` (list[int]): body paragraph indices to target.  Empty = all.
          - ``action`` (str): type of formatting.
          - ``value`` (object): formatting value.

        Returns a list of applied action descriptions.
        """
        applied: list[str] = []
        for act in actions:
            action = act.get("action")
            value = act.get("value")
            indices = act.get("indices", [])
            DocumentParser._apply_format_action(doc, action, value, indices)
            desc = f"{action} → {value}"
            if indices:
                desc += f" (paragraphs {indices})"
            else:
                desc += " (all paragraphs)"
            applied.append(desc)
        return applied

    @staticmethod
    def _apply_format_action(
        doc: DocxDocument,
        action: str,
        value: object,
        target_indices: list[int] | None = None,
    ) -> None:
        """Apply a formatting action to paragraphs in *doc*.

        If *target_indices* is None or empty, applies to ALL paragraphs.
        Otherwise only body paragraphs at those indices are affected.
        """

        # Collect all paragraphs
        all_paras: list = list(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_paras.extend(cell.paragraphs)
        # Text box paragraphs (XML level)
        txbx_paras: list = []
        for txbx in doc.element.findall(f'.//{qn("w:txbxContent")}'):
            txbx_paras.extend(txbx.findall(qn('w:p')))

        # Filter to target indices if specified
        if target_indices:
            filtered_paras = []
            for i in target_indices:
                if i < len(doc.paragraphs):
                    filtered_paras.append(doc.paragraphs[i])
            all_paras = filtered_paras

        if action == "line_spacing":
            is_double = value == "double"
            for para in all_paras:
                DocumentParser.set_line_spacing(para, is_double)
            for p_elem in txbx_paras:
                DocumentParser.set_textbox_line_spacing(p_elem)

        elif action == "font_color":
            color_hex = str(value)
            for para in all_paras:
                for run in para.runs:
                    try:
                        run.font.color.rgb = RGBColor.from_string(color_hex)
                    except (ValueError, AttributeError):
                        pass
            # Text box runs (XML level)
            for p_elem in txbx_paras:
                for r in p_elem.findall(qn('w:r')):
                    rPr = r.find(qn('w:rPr'))
                    if rPr is None:
                        rPr = OxmlElement('w:rPr')
                        r.insert(0, rPr)
                    color_el = rPr.find(qn('w:color'))
                    if color_el is None:
                        color_el = OxmlElement('w:color')
                        rPr.append(color_el)
                    color_el.set(qn('w:val'), color_hex)

        elif action == "font_name":
            font_name = str(value)
            for para in all_paras:
                for run in para.runs:
                    run.font.name = font_name
            # Text box runs (XML level)
            for p_elem in txbx_paras:
                for r in p_elem.findall(qn('w:r')):
                    rPr = r.find(qn('w:rPr'))
                    if rPr is None:
                        rPr = OxmlElement('w:rPr')
                        r.insert(0, rPr)
                    rFonts = rPr.find(qn('w:rFonts'))
                    if rFonts is None:
                        rFonts = OxmlElement('w:rFonts')
                        rPr.append(rFonts)
                    for tag in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
                        rFonts.set(qn(tag), font_name)

        elif action == "bold":
            for para in all_paras:
                for run in para.runs:
                    run.font.bold = bool(value)
            for p_elem in txbx_paras:
                for r in p_elem.findall(qn('w:r')):
                    rPr = r.find(qn('w:rPr'))
                    if bool(value):
                        if rPr is None:
                            rPr = OxmlElement('w:rPr')
                            r.insert(0, rPr)
                        if rPr.find(qn('w:b')) is None:
                            rPr.append(OxmlElement('w:b'))
                    else:
                        b_el = rPr.find(qn('w:b')) if rPr is not None else None
                        if b_el is not None:
                            rPr.remove(b_el)

        elif action == "italic":
            for para in all_paras:
                for run in para.runs:
                    run.font.italic = bool(value)
            for p_elem in txbx_paras:
                for r in p_elem.findall(qn('w:r')):
                    rPr = r.find(qn('w:rPr'))
                    if bool(value):
                        if rPr is None:
                            rPr = OxmlElement('w:rPr')
                            r.insert(0, rPr)
                        if rPr.find(qn('w:i')) is None:
                            rPr.append(OxmlElement('w:i'))
                    else:
                        i_el = rPr.find(qn('w:i')) if rPr is not None else None
                        if i_el is not None:
                            rPr.remove(i_el)

        elif action == "clear_shading":
            DocumentParser.clear_background_shading(doc)
